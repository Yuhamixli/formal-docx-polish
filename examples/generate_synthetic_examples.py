from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = REPO_ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from formal_docx_polish.polish import polish_file  # noqa: E402
from formal_docx_polish.validate import validate_file  # noqa: E402


OUT_DIR = Path(__file__).resolve().parent / "generated"


def _build_request_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("集团财务公司")
    doc.add_paragraph("关于组建数智化部的请示")
    doc.add_paragraph("围绕公司数智化转型与大会筹备工作的设立建议")
    doc.add_paragraph("（可报批版）")
    doc.add_paragraph("2026年3月")
    doc.add_paragraph("主送：公司党委、董事会、经理层")
    doc.add_paragraph("事项：提请审议同意组建数智化部，并授权统筹推进公司数智化体系建设相关工作")
    doc.add_paragraph("一、请示背景")
    doc.add_paragraph("当前，公司内外部发展环境正在发生深刻变化。")
    doc.add_paragraph("二、组建数智化部的必要性")
    doc.add_paragraph("（一）落实战略要求和监管导向。")
    doc.add_paragraph("（二）补齐数据治理和平台统筹能力。")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "职责模块"
    table.rows[0].cells[2].text = "职责内容"
    row = table.add_row()
    row.cells[0].text = "1"
    row.cells[1].text = "数据治理"
    row.cells[2].text = "统一数据标准、口径、质量和血缘"
    doc.save(path)


def _build_plan_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("集团财务公司")
    doc.add_paragraph("数智化部建设总体方案")
    doc.add_paragraph("2026—2028年")
    doc.add_paragraph("（可汇报版）")
    doc.add_paragraph("2026年3月")
    doc.add_paragraph("一、总体判断")
    doc.add_paragraph("公司数智化建设已经进入必须系统化推进的新阶段。")
    doc.add_paragraph("二、建设目标")
    doc.add_paragraph("（一）总体目标")
    doc.add_paragraph("到 2028 年基本建成覆盖主要经营管理场景的数智化能力体系。")
    doc.save(path)


def _build_regulation_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("集团财务公司")
    doc.add_paragraph("数智化部治理制度框架")
    doc.add_paragraph("试行")
    doc.add_paragraph("（讨论版）")
    doc.add_paragraph("2026年3月")
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("第一条 为规范公司数智化项目建设与运行，制定本制度框架。")
    doc.add_paragraph("第二章 治理组织与职责")
    doc.add_paragraph("第二条 数智化部为公司数智化工作的归口管理部门。")
    doc.save(path)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    request_src = OUT_DIR / "request-source.docx"
    request_polished = OUT_DIR / "request-polished.docx"
    plan_src = OUT_DIR / "plan-source.docx"
    regulation_src = OUT_DIR / "regulation-source.docx"
    request_validation = OUT_DIR / "request-validation.json"

    _build_request_doc(request_src)
    _build_plan_doc(plan_src)
    _build_regulation_doc(regulation_src)
    polish_file(request_src, request_polished, document_kind="request")
    validation = validate_file(request_polished, document_kind="request")
    request_validation.write_text(
        json.dumps(validation, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"Generated examples in: {OUT_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
