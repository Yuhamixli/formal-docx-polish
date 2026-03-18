from pathlib import Path

from docx import Document

from formal_docx_polish.polish import polish_file
from formal_docx_polish.validate import validate_file


def test_polish_then_validate(tmp_path: Path) -> None:
    src = tmp_path / "input.docx"
    out = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("集团财务公司")
    doc.add_paragraph("关于组建数智化部的请示")
    doc.add_paragraph("围绕公司数智化转型与大会筹备工作的设立建议")
    doc.add_paragraph("（可报批版）")
    doc.add_paragraph("2026年3月")
    doc.add_paragraph("一、请示背景")
    doc.add_paragraph("当前，公司内外部发展环境正在发生深刻变化。")
    doc.save(src)

    polish_file(src, out, document_kind="request")
    result = validate_file(out, document_kind="request")
    assert result["ok"] is True
