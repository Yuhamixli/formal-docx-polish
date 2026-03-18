from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from formal_docx_polish.profile import load_profile


_DATE_RE = re.compile(r"^\d{4}年\d{1,2}月(?:\d{1,2}日)?$")
_VERSION_RE = re.compile(r"^(（.*(?:版|稿)）|试行|V\d.*)$")
_CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百千]+章")
_ARTICLE_RE = re.compile(r"^第[一二三四五六七八九十百千]+条")
_H1_RE = re.compile(r"^[一二三四五六七八九十百千]+、")
_H2_RE = re.compile(r"^（[一二三四五六七八九十百千]+）")
_H3_RE = re.compile(r"^\d+[\.．、]")


def _paragraph_text(paragraph: Any) -> str:
    return (paragraph.text or "").strip()


def _first_nonempty_run(paragraph: Any) -> Any | None:
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return run
    return None


def _set_run_font(run: Any, spec: dict[str, Any]) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = spec.get("font_latin", "Times New Roman")
    run.font.size = Pt(spec.get("font_size_pt", 16))
    run.bold = bool(spec.get("bold", False))
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), spec.get("font_east_asia", "仿宋_GB2312"))


def _set_paragraph_format(paragraph: Any, spec: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(spec.get("line_spacing_pt", 28))
    pf.space_before = Pt(spec.get("space_before_pt", 0))
    pf.space_after = Pt(spec.get("space_after_pt", 0))
    pf.first_line_indent = Pt(
        spec.get("font_size_pt", 16) * spec.get("first_line_indent_chars", 0),
    )
    paragraph.alignment = alignment_map.get(
        spec.get("alignment", "left"),
        WD_ALIGN_PARAGRAPH.LEFT,
    )


def _apply_paragraph_style(paragraph: Any, spec: dict[str, Any]) -> None:
    _set_paragraph_format(paragraph, spec)
    for run in paragraph.runs:
        if run.text:
            _set_run_font(run, spec)


def find_body_start(document: Any) -> int:
    for idx, paragraph in enumerate(document.paragraphs):
        text = _paragraph_text(paragraph)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name:
            return idx
        if any(
            pattern.match(text)
            for pattern in (_CHAPTER_RE, _ARTICLE_RE, _H1_RE, _H2_RE, _H3_RE)
        ):
            return idx
    return len(document.paragraphs)


def detect_role(text: str) -> str:
    if _CHAPTER_RE.match(text):
        return "chapter"
    if _ARTICLE_RE.match(text):
        return "article"
    if _H1_RE.match(text):
        return "heading1"
    if _H2_RE.match(text):
        return "heading2"
    if _H3_RE.match(text):
        return "heading3"
    return "body"


def _style_cover(document: Any, profile: dict[str, Any], body_start: int) -> None:
    nonempty = [
        paragraph
        for paragraph in document.paragraphs[:body_start]
        if _paragraph_text(paragraph)
    ]
    if not nonempty:
        return

    org_text = _paragraph_text(nonempty[0])
    title_done = False
    for idx, paragraph in enumerate(nonempty):
        text = _paragraph_text(paragraph)
        if idx == 0 or text == org_text:
            spec = profile["cover_org"]
        elif not title_done:
            spec = profile["cover_title"]
            title_done = True
        elif _DATE_RE.match(text):
            spec = profile["cover_date"]
        elif _VERSION_RE.match(text):
            spec = profile["cover_version"]
        else:
            spec = profile["cover_subtitle"]
        _apply_paragraph_style(paragraph, spec)


def _style_body(document: Any, profile: dict[str, Any], body_start: int) -> int:
    styled = 0
    for paragraph in document.paragraphs[body_start:]:
        text = _paragraph_text(paragraph)
        if not text:
            continue
        role = detect_role(text)
        spec = profile.get(role, profile["body"])
        _apply_paragraph_style(paragraph, spec)
        styled += 1
    return styled


def _style_tables(document: Any, profile: dict[str, Any]) -> int:
    from docx.enum.table import WD_ALIGN_VERTICAL

    styled = 0
    for table in document.tables:
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        try:
            table.autofit = False
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    spec = (
                        profile["table_header"]
                        if row_idx == 0
                        else profile["table_body"]
                    )
                    _apply_paragraph_style(paragraph, spec)
                    styled += 1
    return styled


def _ensure_page_numbers(document: Any, profile: dict[str, Any]) -> bool:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = document.sections[0].footer
    existing = "".join(p.text for p in footer.paragraphs).strip()
    if existing:
        return False

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spec = profile["footer_page_number"]

    def _font_run(text: str) -> Any:
        run = para.add_run(text)
        _set_run_font(run, spec)
        return run

    _font_run("- ")
    r_begin = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fld_begin)

    r_instr = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r_instr._r.append(instr)

    r_sep = para.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(fld_sep)

    _font_run("1")
    r_end = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fld_end)
    _font_run(" -")
    return True


def polish_document(document: Any, document_kind: str = "generic") -> dict[str, Any]:
    from docx.shared import Cm

    profile = load_profile(document_kind)
    page = profile["page"]
    for section in document.sections:
        section.page_width = Cm(page["width_cm"])
        section.page_height = Cm(page["height_cm"])
        section.top_margin = Cm(page["top_margin_cm"])
        section.bottom_margin = Cm(page["bottom_margin_cm"])
        section.left_margin = Cm(page["left_margin_cm"])
        section.right_margin = Cm(page["right_margin_cm"])

    body_start = find_body_start(document)
    _style_cover(document, profile, body_start)
    styled_paragraphs = _style_body(document, profile, body_start)
    styled_table_paragraphs = _style_tables(document, profile)
    page_numbers_added = _ensure_page_numbers(document, profile)
    return {
        "document_kind": document_kind,
        "body_start_index": body_start,
        "styled_paragraphs": styled_paragraphs,
        "styled_table_paragraphs": styled_table_paragraphs,
        "page_numbers_added": page_numbers_added,
    }


def polish_file(input_path: str | Path, output_path: str | Path, document_kind: str = "generic") -> dict[str, Any]:
    input_fp = Path(input_path)
    output_fp = Path(output_path)
    document = Document(str(input_fp))
    result = polish_document(document, document_kind=document_kind)
    output_fp.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_fp))
    return {
        "input": str(input_fp),
        "output": str(output_fp),
        **result,
    }
