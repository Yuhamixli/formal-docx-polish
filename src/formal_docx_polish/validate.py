from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from formal_docx_polish.polish import (
    _first_nonempty_run,
    _paragraph_text,
    detect_role,
    find_body_start,
)
from formal_docx_polish.profile import load_profile


def _cm_value(length: Any) -> float | None:
    try:
        return float(length.cm)
    except Exception:
        return None


def _pt_value(length: Any) -> float | None:
    try:
        return float(length.pt)
    except Exception:
        return None


def _approx_equal(left: float | None, right: float | None, tolerance: float) -> bool:
    if left is None or right is None:
        return False
    return abs(left - right) <= tolerance


def validate_document(document: Any, document_kind: str = "generic") -> dict[str, Any]:
    profile = load_profile(document_kind)
    issues: list[str] = []

    page = profile["page"]
    section = document.sections[0]
    margin_checks = {
        "top_margin_cm": (_cm_value(section.top_margin), page["top_margin_cm"]),
        "bottom_margin_cm": (_cm_value(section.bottom_margin), page["bottom_margin_cm"]),
        "left_margin_cm": (_cm_value(section.left_margin), page["left_margin_cm"]),
        "right_margin_cm": (_cm_value(section.right_margin), page["right_margin_cm"]),
    }
    for name, (actual, expected) in margin_checks.items():
        if not _approx_equal(actual, expected, 0.2):
            issues.append(f"{name}: expected {expected}, got {actual}")

    body_start = find_body_start(document)
    cover = [p for p in document.paragraphs[:body_start] if _paragraph_text(p)]
    if len(cover) < 4:
        issues.append("cover appears incomplete before first heading")

    sample_checked = 0
    for paragraph in document.paragraphs[body_start:]:
        text = _paragraph_text(paragraph)
        if not text:
            continue
        role = detect_role(text)
        spec = profile.get(role, profile["body"])
        line_spacing = _pt_value(paragraph.paragraph_format.line_spacing)
        expected_line_spacing = float(spec.get("line_spacing_pt", 28))
        if line_spacing is None or not _approx_equal(line_spacing, expected_line_spacing, 1.5):
            issues.append(
                f"{role} line spacing mismatch: '{text[:24]}' expected {expected_line_spacing}, got {line_spacing}",
            )
            break

        expected_indent = spec.get("font_size_pt", 16) * spec.get("first_line_indent_chars", 0)
        indent = _pt_value(paragraph.paragraph_format.first_line_indent) or 0.0
        if abs(indent - expected_indent) > 2:
            issues.append(
                f"{role} indent mismatch: '{text[:24]}' expected {expected_indent}, got {indent}",
            )
            break

        run = _first_nonempty_run(paragraph)
        if run is not None:
            size = _pt_value(run.font.size) or _pt_value(paragraph.style.font.size)
            expected_size = float(spec.get("font_size_pt", 16))
            if size is None or abs(size - expected_size) > 1:
                issues.append(
                    f"{role} font size mismatch: '{text[:24]}' expected {expected_size}, got {size}",
                )
                break
        sample_checked += 1
        if sample_checked >= 18:
            break

    if document.tables:
        first_cell_para = document.tables[0].rows[0].cells[0].paragraphs[0]
        header_run = _first_nonempty_run(first_cell_para)
        if header_run is not None and not bool(header_run.bold):
            issues.append("table header is not bold in the first table")

    return {
        "ok": not issues,
        "document_kind": document_kind,
        "issues": issues,
        "checked_body_paragraphs": sample_checked,
        "tables": len(document.tables),
    }


def validate_file(input_path: str | Path, document_kind: str = "generic") -> dict[str, Any]:
    input_fp = Path(input_path)
    document = Document(str(input_fp))
    result = validate_document(document, document_kind=document_kind)
    return {
        "file": str(input_fp),
        **result,
    }
